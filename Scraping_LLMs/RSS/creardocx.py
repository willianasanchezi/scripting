from docx import Document
from datetime import datetime

fecha_hora_actual = datetime.now()

def crear_docx(respuesta, tokens_enviados, tokens_recibidos, tokens_totales):

    msg = respuesta
    try:

        doc = Document()

        paragraph = doc.add_paragraph()
        paragraph.add_run(msg)
        
        doc.add_paragraph(f'tokens_enviados: {str(tokens_enviados)}') 
        doc.add_paragraph(f'tokens_recibidos: {str(tokens_recibidos)}')
        doc.add_paragraph(f'tokens_totales: {str(tokens_totales)}')
        doc.add_paragraph(f'Fecha y Hora creación documento {fecha_hora_actual}')
        
        doc.save("resultado.docx")
        return respuesta

    except Exception as e:
        print(f"Error al procesar: {e}")
        return None